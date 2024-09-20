import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import datetime
from docx.shared import Pt, RGBColor
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import os

class expense:
    def __init__ (self, category: str, name: str, 
                  cost: float, incurrence: str):
        self.name = name
        self.category = category
        self.incurrence = incurrence
        try:
            if incurrence == 'daily':
                self.cost = cost * 31
            elif incurrence == 'weekly':
                self.cost = cost * 4.5
            elif incurrence == 'monthly':
                self.cost = cost
            else:
                raise ValueError("Need 'daily', 'weekly', or 'monthly' for incurrence strings.")
            
            if not(self.category and self.name and self.cost and self.incurrence):
                raise ValueError("Need to instantiate all instances.")

        except ValueError as e:
            print(f"An error occurred: {e}")

class expenses:
    def __init__ (self):
        first_line = {
            'category' : [],
            'name' : [],
            'cost' : [],
            'incurrence' : []
        }
        self.df = pd.DataFrame(first_line)
    
    def add_expense(self, user_expense: expense):
        try:
            if not(user_expense.category and user_expense.name and user_expense.cost and user_expense.incurrence):
                raise ValueError("Need to instantiate all instances.")
            else: 
                self.df.loc[len(self.df)] = [user_expense.category, user_expense.name,
                                             user_expense.cost, user_expense.incurrence]
        except ValueError as e:
            print(f"An error occurred: {e}")

class income:
    def __init__ (self, category: str, name: str, 
                  cost: float, incurrence: str):
        self.name = name
        self.category = category
        self.incurrence = incurrence
        try:
            if incurrence == 'daily':
                self.cost = cost * 31
            elif incurrence == 'weekly':
                self.cost = cost * 4.5
            elif incurrence == 'monthly':
                self.cost = cost
            else:
                raise ValueError("Need 'daily', 'weekly', or 'monthly' for incurrence strings.")
            
            if not(self.category and self.name and self.cost and self.incurrence):
                raise ValueError("Need to instantiate all instances.")

        except ValueError as e:
            print(f"An error occurred: {e}")
        

class incomes:
    def __init__ (self):
        first_line = {
            'category' : [],
            'name' : [],
            'cost' : [],
            'incurrence' : []
        }
        self.df = pd.DataFrame(first_line)
    
    def add_expense(self, user_expense: expense):
        try:
            if not(user_expense.category and user_expense.name and user_expense.cost and user_expense.incurrence):
                raise ValueError("Need to instantiate all instances.")
            else: 
                self.df.loc[len(self.df)] = [user_expense.category, user_expense.name,
                                             user_expense.cost, user_expense.incurrence]
        except ValueError as e:
            print(f"An error occurred: {e}")
        

class budget:
    def __init__ (self):
        self.income_table = None
        self.expense_table = None

    def add_table(self, inc_exp): 
        try:
            if isinstance(inc_exp, incomes):
                self.income_table = inc_exp.df
            elif isinstance(inc_exp, expenses):
                self.expense_table = inc_exp.df
            else:
                raise ValueError("Needs to be an expense or income class.")
        except ValueError as e:
            print(f"An error occurred: {e}")

    def transform_df_to_dict(self, df):
        categories = {}
        for category, group in df.groupby('category'):
            items = [(row['name'], row['cost']) for _, row in group.iterrows()]
            categories[category] = items
        return categories
    
    def create_pie_chart(self, data, title, output_file):
        data = data.groupby('category').sum().reset_index()

        plt.figure(figsize=(8, 8))
        plt.pie(data['cost'], labels=data['category'], autopct='%1.1f%%', startangle=140, colors=sns.color_palette("pastel"))
        plt.title(title)
        plt.savefig(output_file)
        plt.close()

    # Create the Word document with incomes and expenses
    def create_word_document_with_incomes_and_expenses(self,incomes, expenses, output_file):
        doc = Document()
        # Add a title with custom font size
        title = doc.add_heading('Financial Report by Category', level=1)
        run = title.runs[0]
        run.font.size = Pt(24)
        
        # Add a horizontal line
        doc.add_paragraph('---------------------------')
        
        # Add Incomes section
        doc.add_heading('Incomes', level=2)
        total_income = 0
        for category, items in incomes.items():
            p = doc.add_paragraph()
            p.add_run(f'Category: {category}\n').bold = True
            for item, amount in items:
                p.add_run(f'  - {item}: ${amount:.2f}\n')
                total_income += amount
            p.add_run(f'Total for {category}: ${sum(amount for _, amount in items):.2f}\n').italic = True
            p.add_run('\n')
        
        p = doc.add_paragraph()
        p.add_run(f'Total Income: ${total_income:.2f}').bold = True
        p.runs[0].font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # Green color
        
        # Add a horizontal line
        doc.add_paragraph('---------------------------')
        
        # Add Expenses section
        doc.add_heading('Expenses', level=2)
        total_expenses = 0
        for category, items in expenses.items():
            p = doc.add_paragraph()
            p.add_run(f'Category: {category}\n').bold = True
            for item, amount in items:
                p.add_run(f'  - {item}: ${amount:.2f}\n')
                total_expenses += amount
            p.add_run(f'Total for {category}: ${sum(amount for _, amount in items):.2f}\n').italic = True
            p.add_run('\n')

        p = doc.add_paragraph()
        p.add_run(f'Total Expenses: ${total_expenses:.2f}').bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red color

        # Add pie charts to the document
        doc.add_heading('Visual Summary', level=2)
        doc.add_paragraph('Income Distribution:')
        doc.add_picture('income_pie_chart.png', width=Inches(4))
        doc.add_paragraph('Expense Distribution:')
        doc.add_picture('expense_pie_chart.png', width=Inches(4))
        
        # Save the document
        doc.save(output_file)
    
    def run_analysis(self):
        try:
            if (isinstance(self.income_table, pd.DataFrame) and isinstance(self.expense_table, pd.DataFrame)):
                # Get the directory of the current script
                current_script_directory = Path(__file__).parent
                
                # Define the "saves" directory
                saves_directory = current_script_directory / "saves"

                # Ensure the "saves" directory exists
                saves_directory.mkdir(parents=True, exist_ok=True)

                # Define the new directory and file name
                new_directory = datetime.date.today().isoformat()
                file_name = 'overview.docx'

                # Create the full path for the new directory
                new_directory_path = saves_directory / new_directory

                # Create the new directory (if it doesn't exist)
                new_directory_path.mkdir(parents=True, exist_ok=True)

                # Define the full path to the file
                file_path = new_directory_path / file_name

                incomes = self.transform_df_to_dict(self.income_table)
                expenses = self.transform_df_to_dict(self.expense_table)

                # Create pie charts for incomes and expenses
                self.create_pie_chart(self.income_table, 'Income Distribution', 'income_pie_chart.png')
                self.create_pie_chart(self.expense_table, 'Expense Distribution', 'expense_pie_chart.png')

                # Create the Word document with incomes and expenses
                self.create_word_document_with_incomes_and_expenses(incomes, expenses, file_path)

                os.remove('income_pie_chart.png')
                os.remove('expense_pie_chart.png')

            else:
                raise ValueError("Tables not properly represented.")
        except ValueError as e:
            print(f"An error occurred: {e}")

if __name__ == "__main__":
    print()